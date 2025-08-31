#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Erzeugt pro Kapitel Markdown + DOCX mit:
- farbigen Überschriften
- fortlaufender Nummerierung
- optionaler Q&A-Einbindung (am Dokumentende)
Input:
  merged_fully_text.txt    # linearer Text inkl. Bild-OCR an Position
  qa.json                  # optional: [{id, question, answer, pdf}...]
Kapitelbildung:
  - nach PDF-Block á 10 Dateien ODER
  - via Marker '--- PDF: XYZ ---' (aus deiner Pipeline)
"""
import re, json
from pathlib import Path

# ---- Konfiguration ----
IN_TXT   = Path("merged_fully_text.txt")     # aus pipeline_qa_text_with_image.py
IN_QA    = Path("qa.json")                   # optional
OUT_DIR  = Path("styled_out")
OUT_DIR.mkdir(exist_ok=True)

# Kapitel-Logik: pro 10 PDFs (Marker aus Pipeline:  '--- PDF: <name> ---')
PDF_MARK = re.compile(r"^---\s*PDF:\s*(.+?)\s*---\s*$", re.MULTILINE)

# Farben (Hex) für DOCX-Styles
COLOR_H1 = "1F4E79"   # Dunkelblau
COLOR_H2 = "006D77"   # Petrol
COLOR_ACCENT = "4F4F4F"

# ---- Hilfen ----
def split_by_pdfs(text):
    """Erzeugt Liste (pdf_order, chunks_by_pdf) anhand der PDF-Marker."""
    parts = PDF_MARK.split(text)
    # parts = ["<vor erstem>", pdf1, text1, pdf2, text2, ...]
    pdf_order, chunks = [], {}
    # starte bei index 1
    for i in range(1, len(parts), 2):
        pdf = parts[i].strip()
        body = parts[i+1]
        pdf_order.append(pdf)
        chunks[pdf] = body.strip()
    return pdf_order, chunks

def group_into_chapters(pdf_order, k=10):
    """Kapitel à k PDFs."""
    chaps = []
    for i in range(0, len(pdf_order), k):
        chaps.append(pdf_order[i:i+k])
    return chaps

def md_escape(s:str)->str:
    return s.replace("<","&lt;").replace(">","&gt;")

# ---- Markdown export ----
def write_markdown(ch_id:int, pdf_names, chunks_by_pdf, qa_items):
    md = []
    md.append(f"# Kapitel {ch_id}\n")
    md.append("> Hinweis: Inhalte sind linear extrahiert. Bildtexte wurden an Position in Fließtext integriert.\n")

    # Nummerierung: Abschnittsüberschrift je PDF
    for idx, pdf in enumerate(pdf_names, 1):
        md.append(f"\n## {ch_id}.{idx} {md_escape(pdf)}\n")
        body = chunks_by_pdf.get(pdf, "").strip()
        if not body:
            md.append("_(kein Inhalt erkannt)_\n")
        else:
            # leichte Normalisierung: doppelte Leerzeilen reduzieren
            body = re.sub(r"\n{3,}", "\n\n", body)
            md.append(body + "\n")

    # Q&A optional anhängen
    if qa_items:
        md.append("\n---\n")
        md.append(f"## Q&A zu Kapitel {ch_id}\n")
        for q in qa_items:
            md.append(f"**Q{q['id']}** ({q.get('pdf','?')}): {q['question']}\n")
            if q.get("answer"):
                md.append(f"*Antwort:* {q['answer']}\n")
    out = OUT_DIR / f"Kapitel_{ch_id:02d}.md"
    out.write_text("\n".join(md), encoding="utf-8")
    return out

# ---- DOCX export mit Farben/Styles ----
def write_docx(ch_id:int, pdf_names, chunks_by_pdf, qa_items):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Styles anpassen (Heading 1 / 2) – Farbe, Größe
    h1 = doc.styles["Heading 1"].font
    h1.size = Pt(18)
    h1.bold = True
    h1.color.rgb = RGBColor.from_string(COLOR_H1)  # python-docx Font color API. Quelle: docs.  [oai_citation:0‡python-docx.readthedocs.io](https://python-docx.readthedocs.io/en/latest/user/text.html?utm_source=chatgpt.com) [oai_citation:1‡Stack Overflow](https://stackoverflow.com/questions/41979095/write-text-in-particular-font-color-in-ms-word-using-python-docx?utm_source=chatgpt.com)

    h2 = doc.styles["Heading 2"].font
    h2.size = Pt(14)
    h2.bold = True
    h2.color.rgb = RGBColor.from_string(COLOR_H2)

    # Titel
    t = doc.add_heading(f"Kapitel {ch_id}", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Inhalt je PDF
    for idx, pdf in enumerate(pdf_names, 1):
        doc.add_heading(f"{ch_id}.{idx} {pdf}", level=2)
        body = chunks_by_pdf.get(pdf, "").strip()
        if not body:
            p = doc.add_paragraph()
            run = p.add_run("(kein Inhalt erkannt)")
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(COLOR_ACCENT)
        else:
            for para in body.split("\n"):
                para = para.rstrip()
                if not para:
                    doc.add_paragraph("")
                else:
                    p = doc.add_paragraph(para)

    # Q&A Abschnitt
    if qa_items:
        doc.add_heading(f"Q&A zu Kapitel {ch_id}", level=2)
        for q in qa_items:
            qpar = doc.add_paragraph()
            rq = qpar.add_run(f"Q{q['id']} ")
            rq.bold = True
            rq.font.color.rgb = RGBColor.from_string(COLOR_H1)
            doc.add_paragraph(f"Frage: {q['question']}")
            if q.get("answer"):
                doc.add_paragraph(f"Antwort: {q['answer']}")

    out = OUT_DIR / f"Kapitel_{ch_id:02d}.docx"
    doc.save(out)
    return out

# ---- Main ----
text = IN_TXT.read_text(encoding="utf-8", errors="ignore")
pdf_order, chunks = split_by_pdfs(text)
chapters = group_into_chapters(pdf_order, k=10)

qa = []
if IN_QA.exists():
    qa = json.loads(IN_QA.read_text(encoding="utf-8", errors="ignore"))

# Q&A kapitelweise filtern nach Quelle (pdf-Namen)
def qa_for_pdfs(pdf_list):
    if not qa: return []
    s = set(pdf_list)
    out = []
    for item in qa:
        if item.get("pdf") in s:
            out.append(item)
    return out

manifest = []
for ch_idx, pdfs in enumerate(chapters, 1):
    md_path  = write_markdown(ch_idx, pdfs, chunks, qa_for_pdfs(pdfs))
    docx_path= write_docx(ch_idx, pdfs, chunks, qa_for_pdfs(pdfs))
    manifest.append({"chapter": ch_idx, "markdown": md_path.as_posix(), "docx": docx_path.as_posix(), "pdfs": pdfs})

(Path(OUT_DIR) / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
print("OK:", (Path(OUT_DIR) / "manifest.json"))
