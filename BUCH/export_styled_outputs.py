#!/usr/bin/env python3  
import re, json  
from pathlib import Path  
from docx import Document  
from docx.shared import Pt, RGBColor  
from docx.enum.text import WD_COLOR_INDEX  

IN_TXT = Path("all_text_linear.txt")  
IN_QA = Path("all_qna.json")  
OUT = Path("styled_out")  
OUT.mkdir(exist_ok=True)  

PDF_MARK = re.compile(r"^--- PDF: (.+?) ---$", re.MULTILINE)  

def split_pdfs(text):  
  parts = PDF_MARK.split(text)  
  return [(parts[i], parts[i+1].strip()) for i in range(1, len(parts), 2)]  

def group(items, k=10):  
  return [items[i:i+k] for i in range(0, len(items), k)]  

text = IN_TXT.read_text(encoding="utf-8", errors="ignore")  
pdfs = split_pdfs(text)  
chapters = group(pdfs, 10)  
qa = json.loads(IN_QA.read_text(encoding="utf-8", errors="ignore")) if IN_QA.exists() else []  

def md_write(path, lines):  
  path.write_text("\n".join(lines), encoding="utf-8")  

for ch_idx, ch in enumerate(chapters, 1):  
  md = [f"# Kapitel {ch_idx}", "", "Hinweis: Bildtexte wurden inline via OCR eingefügt.", ""]  
  doc = Document()  

  h1 = doc.styles["Heading 1"].font  
  h1.size = Pt(18); h1.bold = True; h1.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dunkelblau  
  h2 = doc.styles["Heading 2"].font  
  h2.size = Pt(14); h2.bold = True; h2.color.rgb = RGBColor(0x00, 0x6D, 0x77)  # Petrol  

  doc.add_heading(f"Kapitel {ch_idx}", level=1)  

  pdf_names = [name for name, _ in ch]  
  for idx, (pdf_name, body) in enumerate(ch, 1):  
    md.append(f"## {ch_idx}.{idx} {pdf_name}")  
    md.append(body)  
    md.append("")  

    doc.add_heading(f"{ch_idx}.{idx} {pdf_name}", level=2)  
    for line in body.split("\n"):  
      if not line:  
        doc.add_paragraph("")  
        continue  
      is_codeish = ("=" in line or "+" in line or "-" in line or "*" in line or "/" in line or line.strip().lower().startswith("formel:"))  
      run = doc.add_paragraph().add_run(line)  
      if is_codeish:  
        run.font.name = "Courier New"  
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25  

  md.append("---")  
  md.append(f"## Q&A zu Kapitel {ch_idx}")  
  doc.add_heading("Q&A", level=2)  
  for item in qa:  
    if item["pdf"] in pdf_names:  
      md.append(f"- **Q{item['id']}**: {item['question']}")  
      if item.get("answer"):  
        md.append(f"  Antwort: {item['answer']}")  
      p = doc.add_paragraph()  
      rq = p.add_run(f"Q{item['id']}: {item['question']}")  
      rq.bold = True; rq.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  
      if item.get("answer"):  
        doc.add_paragraph(f"Antwort: {item['answer']}")  
      if item.get("description"):  
        doc.add_paragraph(f"Beschreibung: {item['description']}")  

  md_write(OUT / f"Kapitel_{ch_idx:02d}.md", md)  
  doc.save(OUT / f"Kapitel_{ch_idx:02d}.docx")  

print("Word- und Markdown-Dateien erstellt im styled_out-Ordner.")
